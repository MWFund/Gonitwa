from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Title
title = doc.add_heading('Zestawienie Umiejętności Szachowych', 0)
subtitle = doc.add_paragraph('ROZDZIAŁ I. PIONEK – „ADEPT” (0 – 800)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Poniższe zestawienie prezentuje umiejętności uporządkowane według poziomów trudności (1–9), bazując na materiale szkoleniowym "Młody Wojownik".')
doc.add_page_break()

# Data Structure (Extracted from Source B)
# Format: {Level_Int: [(Topic_Name, Skill_Text), ...]}

data = {
    1: [
        ("Szachownica i jej właściwości", "Potrafi wskazać kolumny, rzędy, prawidłowo nazywa pola i odszukuje je na szachownicy"),
        ("Poznajemy gońca", "Potrafi prawidłowo wykonać ruch gońcem, zna wszystkie zasady poruszania się tą figurą"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać jedno pole, na którym stał uciekający goniec przeciwnika"),
        ("Poznajemy wieżę", "Potrafi prawidłowo wykonać ruch wieżą, zna wszystkie zasady poruszania się tą figurą"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać jedno pole, na którym stała uciekająca wieża przeciwnika"),
        ("Poznajemy hetmana", "Potrafi prawidłowo wykonać ruch hetmanem, zna wszystkie zasady poruszania się tą figurą"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać jedno pole, na którym stał hetman przeciwnika"),
        ("Poznajemy skoczka", "Potrafi prawidłowo wykonać ruch skoczkiem, zna wszystkie zasady poruszania się tą figurą"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać jedno pole, na którym stał skoczek przeciwnika"),
        ("Poznajemy ruch królem", "Potrafi prawidłowo wykonać ruch królem, zna wszystkie zasady poruszania się tą figurą"),
        ("Poznajemy ruch piona", "Potrafi prawidłowo wykonać ruch pionem, zna zasady poruszania się tą bierką"),
        ("Prawidłowe ustawienie bierek", "Potrafi prawidłowo ustawić bierki na szachownicy, zna rzędy, na których powinny stać białe i czarne"),
        ("Wartość materialna figur", "Zna wartość poszczególnych bierek"),
        ("Pojęcie szacha i obrony", "Zna czym jest szach, potrafi wskazać, która bierka szachuje króla"),
        ("Zakończenie partii - mat", "Zna pojęcie mata"),
        ("Zakończenie partii - mat", "Wybiera właściwą figurę, którą należy wstawić na zaznaczone pole, tak aby był mat"),
        ("Zakończenie partii - mat", "Potrafi ustawić na szachownicy mata mając do dyspozycji określone bierki – własna kompozycja"),
        ("Zakończenie partii - pat", "Zna pojęcie pata – potrafi określić czy na szachownicy jest pat czy mat"),
        ("Zakończenie partii - pat", "Potrafi ustawić pata wstawiając figurę na zaznaczone pole"),
        ("Matowanie hetmanem i wieżą", "Zna dwa schematy matowania króla przeciwnika (tzw. „schemat drabinkowy” lub „schemat schodkowy”)"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla przeciwnika na skrajnej linii bez wsparcia własnego króla"),
        ("Matowanie dwiema wieżami", "Zna schemat matowania króla przeciwnika tzw. „schemat schodkowy”"),
        ("Matowanie dwiema wieżami", "Potrafi zamatować króla przeciwnika na skrajnej linii bez wsparcia własnego króla"),
        ("Ruch specjalny – roszada", "Zna pojęcie roszady, wie kiedy jest niedozwolona i kiedy przejściowo niemożliwa"),
        ("Ruch specjalny – roszada", "Potrafi wykonać krótką i długą roszadę"),
    ],
    2: [
        ("Szachownica i jej właściwości", "Potrafi wskazać przekątne na szachownicy, wymienia kolory pól z pamięci"),
        ("Poznajemy gońca", "Potrafi wykonać bicie gońcem – łamigłówki na 4 ruchy"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać dwa pola, na których stał uciekający goniec"),
        ("Poznajemy wieżę", "Potrafi wykonać bicie wieżą – łamigłówki na 4 ruchy"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać dwa pola, na których stała uciekająca wieża"),
        ("Poznajemy hetmana", "Potrafi wykonać bicie hetmanem – łamigłówki na 4 ruchy"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać dwa pola, na których stał uciekający hetman"),
        ("Poznajemy skoczka", "Potrafi wykonać bicie skoczkiem – łamigłówki na 4 ruchy"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać dwa pola, na których stał uciekający skoczek"),
        ("Poznajemy skoczka", "Rozwiązuje łamigłówki: „problem skoczka szachowego” (wypełnia 40–52 pola)"),
        ("Poznajemy skoczka", "Troyis – przechodzi drugi poziom"),
        ("Poznajemy ruch królem", "Potrafi wykonać bicie królem"),
        ("Poznajemy ruch piona", "Potrafi wykonać bicie pionem – łamigłówki"),
        ("Poznajemy ruch piona", "Zna zasady promocji piona"),
        ("Poznajemy ruch piona", "Potrafi rozegrać i wygrać partię 2 x 2"),
        ("Prawidłowe ustawienie bierek", "Potrafi znaleźć i wskazać brakujące bierki na szachownicy"),
        ("Wartość materialna figur", "Potrafi oszacować wartość dwóch bierek i wskazać różnicę"),
        ("Wartość materialna figur", "Potrafi wskazać, która wymiana jest korzystna"),
        ("Pojęcie szacha i obrony", "Potrafi wykonać szacha"),
        ("Pojęcie szacha i obrony", "Zna wszystkie sposoby obrony przed szachem"),
        ("Zakończenie partii - mat", "Potrafi dać mata w jednym posunięciu przy wskazaniu bierki, którą należy ruszyć"),
        ("Zakończenie partii - pat", "Potrafi ustawić własną kompozycję patową"),
        ("Zakończenie partii - pat", "Potrafi wykonać ruch figurą, który patuje"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla z dowolnej pozycji w mniej niż 50 posunięć"),
        ("Matowanie dwiema wieżami", "Potrafi zamatować króla z dowolnej pozycji w mniej niż 15 posunięć"),
        ("Ruch specjalny – roszada", "Potrafi wskazać którą roszadę można wykonać – łatwe ćwiczenia z jednym utrudnieniem"),
    ],
    3: [
        ("Szachownica i jej właściwości", "Potrafi wskazać 5 pól na szachownicy w ciągu 30 sekund"),
        ("Szachownica i jej właściwości", "Potrafi wymienić wszystkie pola na liniach pionowych podając ich nazwy"),
        ("Poznajemy gońca", "Potrafi wykonać bicie gońcem – łamigłówki na 5 ruchów"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać trzy pola"),
        ("Poznajemy wieżę", "Potrafi wykonać bicie wieżą – łamigłówki na 5 ruchów"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać trzy pola"),
        ("Poznajemy hetmana", "Potrafi wykonać bicie hetmanem – łamigłówki na 5 ruchów"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać trzy pola"),
        ("Poznajemy skoczka", "Potrafi wykonać bicie skoczkiem – łamigłówki na 5 ruchów"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać trzy pola"),
        ("Poznajemy skoczka", "„Problem skoczka szachowego”: wypełnia 50–53 pola"),
        ("Poznajemy skoczka", "Troyis – przechodzi trzeci poziom"),
        ("Poznajemy ruch królem", "Potrafi wskazać pola, na które nie może się poruszyć (pola atakowane przez jedną bierkę przeciwnika)"),
        ("Poznajemy ruch królem", "Zna zasadę: „król do króla się nie przytula” – wskazuje właściwe pola"),
        ("Poznajemy ruch piona", "Zna bicie w przelocie *en passant* i potrafi je zastosować w praktyce"),
        ("Poznajemy ruch piona", "Potrafi rozegrać i wygrać partię 5 x 5"),
        ("Prawidłowe ustawienie bierek", "Wskazuje błędy w ustawieniu szachownicy"),
        ("Wartość materialna figur", "Prawidłowo wykonuje działania na trzech bierkach"),
        ("Pojęcie szacha i obrony", "Potrafi wskazywać ostatni ruch bierek, które zaszachowały króla"),
        ("Pojęcie szacha i obrony", "Potrafi wskazać wszystkie ruchy bierek które w konkretnej pozycji mogą dać szacha"),
        ("Zakończenie partii - mat", "Potrafi dać mata w jednym posunięciu przy małej ilości bierek na szachownicy"),
        ("Zakończenie partii - pat", "Potrafi znaleźć jedno posunięcie, którym ratuje się od porażki – wymusza pata"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla z dowolnej pozycji w mniej niż 30 posunięć"),
        ("Matowanie dwiema wieżami", "Potrafi zamatować króla z dowolnej pozycji w mniej niż 10 posunięć"),
        ("Ruch specjalny – roszada", "Potrafi wskazać którą roszadę można wykonać – ćwiczenia z większą ilością utrudnień"),
    ],
    4: [
        ("Szachownica i jej właściwości", "Potrafi wskazać 7 pól na szachownicy w ciągu 30 sekund"),
        ("Poznajemy gońca", "Potrafi wykonać bicie gońcem – łamigłówki na 6 ruchów"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać cztery pola"),
        ("Poznajemy wieżę", "Potrafi wykonać bicie wieżą – łamigłówki na 6 ruchów"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać cztery pola"),
        ("Poznajemy hetmana", "Potrafi wykonać bicie hetmanem – łamigłówki na 6 ruchów"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać cztery pola"),
        ("Poznajemy skoczka", "Potrafi wykonać bicie skoczkiem – łamigłówki na 6 ruchów"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać cztery pola"),
        ("Poznajemy skoczka", "„Problem skoczka szachowego”: wypełnia 54–55 pola"),
        ("Poznajemy skoczka", "Troyis – przechodzi czwarty poziom"),
        ("Poznajemy ruch królem", "Potrafi wskazać pola atakowane (gdy są dwie bierki przeciwnika)"),
        ("Prawidłowe ustawienie bierek", "Wskazuje z pamięci kolory pól, na których znajdują się poszczególne bierki"),
        ("Wartość materialna figur", "Prawidłowo wykonuje działania na czterech bierkach"),
        ("Pojęcie szacha i obrony", "Znajduje wszystkie możliwe obrony przed szachem"),
        ("Zakończenie partii - mat", "Potrafi dać mata w jednym posunięciu przy nieograniczonej ilości bierek na szachownicy"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla z dowolnej pozycji w mniej niż 15 posunięć"),
        ("Matowanie dwiema wieżami", "Potrafi zamatować króla przeciwnika przy wsparciu własnego króla na liniach innych niż skrajnej"),
        ("Ruch specjalny – roszada", "Rozwiązuje proste łamigłówki związane z roszadą"),
    ],
    5: [
        ("Szachownica i jej właściwości", "Potrafi wskazać 9 pól na szachownicy w ciągu 30 sekund"),
        ("Szachownica i jej właściwości", "Potrafi wskazać przekątne na szachownicy oraz wymienić nazwy jej pól"),
        ("Poznajemy gońca", "Potrafi wykonać bicie gońcem – łamigłówki na 7 ruchów"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać pięć pól"),
        ("Poznajemy wieżę", "Potrafi wykonać bicie wieżą – łamigłówki na 7 ruchów"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać pięć pól"),
        ("Poznajemy hetmana", "Potrafi wykonać bicie hetmanem – łamigłówki na 7 ruchów"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać pięć pól"),
        ("Poznajemy skoczka", "Potrafi wykonać bicie skoczkiem – łamigłówki na 7 ruchów"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać pięć pól"),
        ("Poznajemy skoczka", "„Problem skoczka szachowego”: wypełnia 56–57 pola"),
        ("Poznajemy skoczka", "Troyis – przechodzi siódmy poziom"),
        ("Poznajemy ruch królem", "Potrafi wskazać pola atakowane (gdy są trzy bierki przeciwnika)"),
        ("Wartość materialna figur", "Prawidłowo wykonuje działania na pięciu bierkach"),
        ("Pojęcie szacha i obrony", "Wskazuje wszystkie szachy jakie można było zrobić w partii (ćwiczenia z partią)"),
        ("Zakończenie partii - mat", "Rozwiązuje proste forsowne maty w dwóch posunięciach przy małej ilości bierek"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla z dowolnej pozycji w mniej niż 10 posunięć"),
        ("Matowanie dwiema wieżami", "Potrafi zamatować króla na linii wskazanej przez przeciwnika"),
        ("Ruch specjalny – roszada", "Rozwiązuje trudniejsze łamigłówki związane z roszadą"),
    ],
    6: [
        ("Szachownica i jej właściwości", "Potrafi wskazać 11 pól na szachownicy w ciągu 30 sekund"),
        ("Szachownica i jej właściwości", "Potrafi z pamięci wymienić narożne pola na szachownicy oraz pamięta kolory tych pól"),
        ("Poznajemy gońca", "Potrafi wykonać bicie gońcem – łamigłówki na 8 ruchów"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać sześć pól"),
        ("Poznajemy wieżę", "Potrafi wykonać bicie wieżą – łamigłówki na 8 ruchów"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać sześć pól"),
        ("Poznajemy hetmana", "Potrafi wykonać bicie hetmanem – łamigłówki na 8 ruchów"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać sześć pól"),
        ("Poznajemy skoczka", "Potrafi wykonać bicie skoczkiem – łamigłówki na 8 ruchów"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać sześć pól"),
        ("Poznajemy skoczka", "„Problem skoczka szachowego”: wypełnia 58–59 pola"),
        ("Poznajemy skoczka", "Troyis – przechodzi dziewiąty poziom"),
        ("Pojęcie szacha i obrony", "Wskazuje wszystkie szachy jakie można było zrobić w partii i wskazuje wszystkie możliwe obrony"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla przeciwnika przy wsparciu własnego króla na liniach innych niż skrajnej"),
        ("Matowanie dwiema wieżami", "Potrafi zamatować króla na polu wskazanym przez przeciwnika"),
    ],
    7: [
        ("Szachownica i jej właściwości", "Potrafi wskazać 13 pól na szachownicy w ciągu 30 sekund"),
        ("Szachownica i jej właściwości", "Potrafi z pamięci wymienić pola na liniach i podać kolory pól na pierwszej i ostatniej linii"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać siedem pól"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać siedem pól"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać siedem pól"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać siedem pól"),
        ("Poznajemy skoczka", "„Problem skoczka szachowego”: wypełnia 60–61 pola"),
        ("Poznajemy skoczka", "Troyis – przechodzi jedenasty poziom"),
        ("Wartość materialna figur", "Prawidłowo wykonuje działania na sześciu bierkach"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla na linii wskazanej przez przeciwnika"),
    ],
    8: [
        ("Szachownica i jej właściwości", "Potrafi z pamięci wskazać kolor każdego pola na szachownicy"),
        ("Szachownica i jej właściwości", "Potrafi z pamięci wymienić wszystkie pola na przekątnych"),
        ("Szachownica i jej właściwości", "Potrafi z pamięci wymienić wszystkie pola na liniach i podać kolory tych pól"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać osiem pól"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać osiem pól"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać osiem pól"),
        ("Poznajemy hetmana", "Rozwiązuje łamigłówkę: „problem ośmiu hetmanów”"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać osiem pól"),
        ("Poznajemy skoczka", "„Problem skoczka szachowego”: wypełnia 62–63 pola"),
        ("Poznajemy skoczka", "Troyis – przechodzi trzynasty poziom"),
        ("Matowanie hetmanem i wieżą", "Potrafi zamatować króla na polu wskazanym przez przeciwnika"),
    ],
    9: [
        ("Szachownica i jej właściwości", "Potrafi wskazać 15 pól na szachownicy w ciągu 30 sekund"),
        ("Poznajemy gońca", "Gra „Uciekająca figura”: potrafi zapamiętać dziewięć pól"),
        ("Poznajemy wieżę", "Gra „Uciekająca figura”: potrafi zapamiętać dziewięć pól"),
        ("Poznajemy hetmana", "Gra „Uciekająca figura”: potrafi zapamiętać dziewięć pól"),
        ("Poznajemy hetmana", "Rozwiązuje łamigłówkę: „problem pięciu hetmanów”"),
        ("Poznajemy skoczka", "Gra „Uciekająca figura”: potrafi zapamiętać dziewięć pól"),
        ("Poznajemy skoczka", "„Problem skoczka szachowego”: wypełnia 64 pola"),
        ("Poznajemy skoczka", "Troyis – przechodzi piętnasty poziom"),
        ("Wartość materialna figur", "Prawidłowo wykonuje działania na siedmiu bierkach"),
        ("Zakończenie partii - pat", "Potrafi znaleźć dwa posunięcia, którymi ratuje się od porażki – wymusza pata"),
    ]
}

# Generate Content
for level in range(1, 10):
    # Header for Level
    h = doc.add_heading(f'POZIOM {level}', level=1)
    
    # Check if we have data for this level
    if level in data and data[level]:
        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Temat'
        hdr_cells[1].text = 'Wymagania / Umiejętności'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        # Fill data
        for topic, skill in data[level]:
            row_cells = table.add_row().cells
            row_cells[0].text = topic
            row_cells[1].text = skill
            
    else:
        doc.add_paragraph('Brak zdefiniowanych umiejętności dla tego poziomu w materiale źródłowym.')
        
    doc.add_paragraph('') # Spacing

# Save
file_path = r"E:\kilo-api\chess_exam_Doc.docx"
doc.save(file_path)

file_path